from flask import Flask, request, jsonify
from flask_cors import CORS
import psycopg2
from psycopg2.extras import RealDictCursor
from werkzeug.utils import secure_filename
from supabase import create_client, Client
import tempfile
import os
from io import BytesIO
import docx
from pinecone import Pinecone, ServerlessSpec
from langchain_community.document_loaders import TextLoader, PyPDFLoader, Docx2txtLoader
from langchain_text_splitters import CharacterTextSplitter
from langchain_pinecone import PineconeVectorStore
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.schema import SystemMessage, HumanMessage
from langchain_community.document_loaders import FireCrawlLoader
from langchain.schema import Document

app = Flask(__name__)
CORS(app)

    
@app.route("/api/search_documents", methods=['POST'])
def search_documents():
    try:
        data = request.get_json()
        query = data.get('text')
        
        if not query:
            return jsonify({'error': 'No query provided'}), 400
        
        pc = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))

        embeddings = OpenAIEmbeddings(
            api_key=os.environ.get('OPENAI_API_KEY'),
            model="text-embedding-3-small"
        )
        vector_store = PineconeVectorStore(
            index_name="vetai-docs",
            embedding=embeddings
        )
        results = vector_store.similarity_search_with_score(
            query=query,
            k=3
        )
        docs = [result[0] for result in results]
        combined_input = (
            "Here are some documents that might help answer the question: "
            + query
            + "\n\nRelevant Documents:\n"
            + "\n\n".join([doc.page_content for doc in docs])
            + "\n\nPlease provide an answer based only on the provided documents. If the answer is not found in the documents, respond with 'I'm not sure'."
        )
        model = ChatOpenAI(
            api_key=os.environ.get('OPENAI_API_KEY'),
            model="gpt-4"
        )
        messages = [
            SystemMessage(content="You are a helpful veterinarian assistant."),
            HumanMessage(content=combined_input),
        ]
        result = model.invoke(messages)
        
        return jsonify({
            'results': result.content,
            'sources': [{'content': doc.page_content, 'metadata': doc.metadata} for doc in docs]
        }), 200
        
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
@app.route("/api/get_files", methods=['GET'])
def get_files():
    try:
        url = os.environ.get('SUPABASE_URL')
        key = os.environ.get('SUPABASE_KEY')
        if not url or not key:
            return jsonify({'error': 'Supabase URL or Key is missing'}), 500
        supabase = create_client(url, key)
        response = supabase.table("files").select("*").execute()
        return jsonify({'files': response.data}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route("/api/delete_file", methods=['POST'])
def delete_files():
    try:
        data = request.get_json()
        if not data or 'file_name' not in data:
            return jsonify({'error': 'No file part'}), 400
        file_name = data['file_name']

        pc = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))
        index = pc.Index("vetai-docs")
        results = index.query(
            vector=[0] * 1536,
            filter={
                "filename": {"$eq": file_name}
            },
            top_k=10000,
            include_metadata=True
        )
        if hasattr(results, 'matches') and results.matches:
            ids_to_delete = [match.id for match in results.matches]
            try:
                index.delete(
                    ids=ids_to_delete
                )
            except Exception as e:
                raise
            
        url = os.environ.get('SUPABASE_URL')
        key = os.environ.get('SUPABASE_KEY')
        supabase = create_client(url, key)
        supabase.table("files").delete().eq("file_name", file_name).execute()

        return jsonify({'message': 'File deleted successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route("/api/upload_file", methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        file = request.files['file']
        file_content = file.read()

        filename = secure_filename(file.filename)
        file_size = int(request.form.get('file_size', 0))
        pc = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))
        embeddings = OpenAIEmbeddings(
            api_key=os.environ.get('OPENAI_API_KEY'),
            model="text-embedding-3-small"
        )
        documents = None
        if file.filename.endswith('.pdf'):
            loader = PyPDFLoader(BytesIO(file_content))
            documents = loader.load()
        elif file.filename.endswith('.docx'):
            doc = docx.Document(BytesIO(file_content))
            text = "\n".join([p.text for p in doc.paragraphs])
            documents = [Document(page_content=text, metadata={})]
        else:
            text = file_content.decode('utf-8')
            documents = [Document(page_content=text, metadata={})]

        text_splitter = CharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        texts = text_splitter.split_documents(documents)
        for text in texts:
            text.metadata["filename"] = filename
        PineconeVectorStore.from_documents(
            documents=texts,
            embedding=embeddings,
            index_name="vetai-docs",
        )

        sup_url = os.environ.get('SUPABASE_URL')
        key = os.environ.get('SUPABASE_KEY')
        supabase: Client = create_client(sup_url, key)
        response = supabase.table("files").insert(
            {"file_name": filename, "file_size": file_size}).execute()

        return jsonify({'message': 'File processed and embeddings stored successfully'}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route("/api/upload_website", methods=['POST'])
def upload_website():
    try:
        api_key = os.environ.get('FIRECRAWL_API_KEY')
        url = request.get_json().get('url')
        loader = FireCrawlLoader(api_key=api_key, url=url, mode="scrape")
        docs = loader.load()
        pc = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))
        text_splitter = CharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
            )
        texts = text_splitter.split_documents(docs)

        for text in texts:
            text.metadata["filename"] = url

        embeddings = OpenAIEmbeddings(
            api_key=os.environ.get('OPENAI_API_KEY'),
            model="text-embedding-3-small"
        )

        PineconeVectorStore.from_documents(
            documents=texts,
            embedding=embeddings,
            index_name="vetai-docs",
        )
        total_bytes = sum(len(doc.page_content.encode('utf-8')) for doc in docs)
        sup_url = os.environ.get('SUPABASE_URL')
        key = os.environ.get('SUPABASE_KEY')

        supabase: Client = create_client(sup_url, key)
        response = supabase.table("files").insert({"file_name": url, "file_size": total_bytes}).execute()
        return jsonify({'message': 'Website processed and embeddings stored successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route("/api/save_template", methods=['POST'])
def save_template():
    try:
        data = request.get_json()
        name = data.get('name')
        text = data.get('text')

        url = os.environ.get('SUPABASE_URL')
        key = os.environ.get('SUPABASE_KEY')

        supabase: Client = create_client(url, key)
        response = (
            supabase.table("templates")
            .insert({"template_name": name, "content": text})
            .execute()
        )
        return jsonify({'message': 'Template saved successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route("/api/load_templates", methods=['GET'])
def load_templates():
    try:
        url = os.environ.get('SUPABASE_URL')
        key = os.environ.get('SUPABASE_KEY')

        supabase: Client = create_client(url, key)
        response = supabase.table("templates").select("*").execute()

        return jsonify({'templates': response.data}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    

if __name__ == '__main__' and os.environ.get('VERCEL_ENV') != 'production':
    app.run(port=5329, debug=True)
